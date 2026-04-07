from docx import Document
import uuid
import os


def generate_doc(text):
    doc = Document()
    doc.add_heading("Documento generado", 0)
    doc.add_paragraph(text)

    os.makedirs("outputs", exist_ok=True)

    file_name = f"{uuid.uuid4()}.docx"
    path = f"outputs/{file_name}"

    doc.save(path)
    return path
